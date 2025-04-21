"""Document processing module for DocImprover."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING
from openai import OpenAI
from .config.config import get_settings
from .config.formatting_config import DOCUMENT_FORMATTING
import logging
import re
import os
import tempfile
import shutil
from pathlib import Path
from typing import Dict, Any, Optional
import uuid
from docx.shared import Inches
import io
from PIL import Image
from docx.oxml import parse_xml
from docx.oxml.ns import qn

logger = logging.getLogger('docimprover')

class DocumentProcessor:
    """Process documents using OpenAI's GPT models."""

    # XML namespaces used in Word documents
    NAMESPACES = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'v': 'urn:schemas-microsoft-com:vml'
    }

    def __init__(self):
        """Initialize the document processor."""
        settings = get_settings()
        self.client = OpenAI(
            api_key=settings.openai_api_key,
            base_url="https://api.openai.com/v1"
        )
        self.model = settings.model_name
        self.system_prompt = settings.system_prompt
        self._temp_dir = None
        self._image_map = {}

    def __enter__(self):
        """Context manager entry."""
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        """Context manager exit."""
        self._cleanup_temp_dir()

    def _create_temp_dir(self) -> str:
        """Create a temporary directory for storing images."""
        if self._temp_dir is None:
            self._temp_dir = tempfile.mkdtemp(prefix='docimprover_')
        return self._temp_dir

    def _cleanup_temp_dir(self) -> None:
        """Clean up temporary directory and reset image map."""
        if self._temp_dir and os.path.exists(self._temp_dir):
            try:
                shutil.rmtree(self._temp_dir, ignore_errors=True)
            except Exception as e:
                logger.error(f"Error cleaning up temporary directory: {e}")
        self._temp_dir = None
        self._image_map = {}

    def _extract_image_from_run(self, run, namespaces: Dict[str, str]) -> tuple[bool, str | None, Any | None]:
        """Extract image information from a run."""
        has_image = False
        rid = None
        image_element = None

        # Check for different types of images in order of most common to least common
        image_paths = [
            ('.//wp:inline//a:blip', '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'),
            ('.//pic:pic//a:blip', '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'),
            ('.//w:drawing//pic:pic//a:blip', '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'),
            ('.//v:shape//v:imagedata', '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}href')
        ]

        for xpath, rid_attr in image_paths:
            elements = run._element.findall(xpath, namespaces)
            if elements:
                has_image = True
                image_element = elements[0]
                rid = image_element.get(rid_attr)
                break

        return has_image, rid, image_element

    def extract_text_and_images(self, doc: Document) -> str:
        """Extract text and images from a Word document, replacing images with placeholders."""
        full_text = []
        self._image_map = {}

        # Set maximum image size (10MB)
        MAX_IMAGE_SIZE = 10 * 1024 * 1024  # 10MB in bytes

        for paragraph in doc.paragraphs:
            modified_text = ""
            for run in paragraph.runs:
                has_image, rid, _ = self._extract_image_from_run(run, self.NAMESPACES)
                
                if has_image and rid:
                    try:
                        # Extract image
                        image_part = doc.part.related_parts[rid]
                        image_bytes = image_part.blob
                        
                        # Check image size
                        if len(image_bytes) > MAX_IMAGE_SIZE:
                            logger.warning(f"Image size exceeds maximum allowed size ({MAX_IMAGE_SIZE} bytes). Skipping.")
                            image_id = "This image is too large to process"
                            modified_text += f"[IMAGE: {image_id}]"
                            continue
                        
                        image_id = str(uuid.uuid4())
                        placeholder = f"[IMAGE:{image_id}]"
                        
                        # Save image info
                        self._image_map[image_id] = {
                            'bytes': image_bytes,
                            'content_type': image_part.content_type,
                            'width': None,  # Will be set when scaling
                            'height': None  # Will be set when scaling
                        }
                        
                        modified_text += placeholder
                    except Exception as e:
                        logger.error(f"Error extracting image: {e}")
                        modified_text += run.text
                else:
                    modified_text += run.text
            
            if modified_text.strip():
                full_text.append(modified_text)

        return "\n".join(full_text)

    def _parse_inline_markdown(self, text: str) -> list[tuple[str, dict]]:
        """Parse inline markdown formatting (bold, italic, etc).
        Returns list of (text, format_dict) tuples."""
        formats = []
        current_pos = 0
        
        # Find all markdown patterns
        patterns = [
            (r'\*\*(.+?)\*\*', {'bold': True}),  # Bold **text**
            (r'__(.+?)__', {'bold': True}),      # Bold __text__
            (r'\*(.+?)\*', {'italic': True}),    # Italic *text*
            (r'_(.+?)_', {'italic': True}),      # Italic _text_
            (r'`(.+?)`', {'code': True}),        # Code `text`
        ]
        
        while current_pos < len(text):
            earliest_match = None
            earliest_pos = len(text)
            match_format = {}
            match_content = ""
            
            # Find the earliest match of any pattern
            for pattern, format_dict in patterns:
                match = re.search(pattern, text[current_pos:])
                if match and match.start() + current_pos < earliest_pos:
                    earliest_match = match
                    earliest_pos = match.start() + current_pos
                    match_format = format_dict
                    match_content = match.group(1)
            
            if earliest_match:
                # Add any text before the match
                if earliest_pos > current_pos:
                    formats.append((text[current_pos:earliest_pos], {}))
                # Add the matched text with its formatting
                formats.append((match_content, match_format))
                current_pos = earliest_pos + len(earliest_match.group(0))
            else:
                # Add remaining text
                formats.append((text[current_pos:], {}))
                break
        
        return formats if formats else [(text, {})]

    def _parse_markdown_line(self, line: str) -> tuple[str, str, int]:
        """Parse a markdown line to determine its type and content.
        Returns: (content, style, list_level)
        """
        # Check for headers
        header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if header_match:
            level = len(header_match.group(1))
            return header_match.group(2), f'Heading {level}', 0

        # Check for unordered lists
        list_match = re.match(r'^(\s*)[*+-]\s+(.+)$', line)
        if list_match:
            indent = len(list_match.group(1))
            level = indent // 2
            return list_match.group(2), 'List Bullet', level

        # Check for ordered lists
        ordered_list_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ordered_list_match:
            indent = len(ordered_list_match.group(1))
            level = indent // 2
            return ordered_list_match.group(2), 'List Number', level

        # Check for blockquotes
        quote_match = re.match(r'^\s*>\s+(.+)$', line)
        if quote_match:
            return quote_match.group(1), 'Quote', 0

        # Default case - normal paragraph
        return line, 'Normal', 0

    def _add_formatted_text(self, paragraph, text: str) -> None:
        """Add text to a paragraph with inline formatting."""
        for content, format_dict in self._parse_inline_markdown(text):
            run = paragraph.add_run(content)
            if format_dict.get('bold'):
                run.bold = True
            if format_dict.get('italic'):
                run.italic = True
            if format_dict.get('code'):
                run.font.name = 'Courier New'

    def _apply_formatting(self, paragraph, style_name: str = 'Normal', level: int = 0) -> None:
        """Apply formatting to a paragraph."""
        # Apply the style
        try:
            paragraph.style = style_name
        except ValueError:
            # If style doesn't exist, fall back to normal
            paragraph.style = 'Normal'

        # Get formatting settings
        settings = DOCUMENT_FORMATTING
        format_settings = settings['default_text']
        
        if style_name.startswith('Heading '):
            level_num = int(style_name.split()[-1])
            if level_num in settings['headings']:
                format_settings = settings['headings'][level_num]
        
        # Apply indentation for lists
        if style_name in ['List Bullet', 'List Number']:
            paragraph.paragraph_format.left_indent = Pt(level * 18)  # 18 points per level
            if level > 0:
                paragraph.paragraph_format.first_line_indent = Pt(-18)  # Hanging indent for bullets/numbers

        self._apply_format_settings(paragraph, format_settings)

    def _add_image_to_doc(self, doc: Document, image_id: str) -> Optional[str]:
        """Add an image to the document from the image map."""
        try:
            if image_id not in self._image_map:
                if image_id.strip() == "This image is too large to process":
                    # This is a placeholder for a large image, add it as text
                    p = doc.add_paragraph()
                    p.add_run(f"[IMAGE: {image_id}]")
                    return None
                return None

            image_data = self._image_map[image_id]
            image_bytes = image_data['bytes']
            content_type = image_data['content_type']

            # Create a stream from the image bytes
            image_stream = io.BytesIO(image_bytes)
            
            # Calculate image size based on original dimensions but limit to max dimensions
            MAX_WIDTH_INCHES = 6.0
            MAX_HEIGHT_INCHES = 8.0
            
            with Image.open(image_stream) as img:
                width, height = img.size
                # Reset stream position after reading
                image_stream.seek(0)
                
                # Convert pixels to inches (assuming 96 DPI)
                width_inches = width / 96
                height_inches = height / 96
                
                # Scale down if necessary while maintaining aspect ratio
                if width_inches > MAX_WIDTH_INCHES:
                    scale = MAX_WIDTH_INCHES / width_inches
                    width_inches = MAX_WIDTH_INCHES
                    height_inches *= scale
                
                if height_inches > MAX_HEIGHT_INCHES:
                    scale = MAX_HEIGHT_INCHES / height_inches
                    height_inches = MAX_HEIGHT_INCHES
                    width_inches *= scale

            # Add a paragraph to contain the image
            p = doc.add_paragraph()
            run = p.add_run()
            
            # Add the image with calculated dimensions
            picture = run.add_picture(image_stream, width=Inches(width_inches))
            
            # Ensure the image is properly embedded in the document relationships
            image_part = doc.part.get_or_create_image(image_bytes)
            rId = doc.part.relate_to(image_part, qn('http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'))
            picture._inline.graphic.graphicData.pic.blipFill.blip.embed = rId
            
            # Clear the stream
            image_stream.close()
            
            return None  # No error
        except Exception as e:
            error_msg = f"Error adding image {image_id}: {str(e)}"
            logger.error(error_msg)
            return error_msg

    def improve_document(self, doc: Document) -> Dict[str, Any]:
        """Improve a document using OpenAI's GPT model."""
        try:
            # Extract text and images
            original_text = self.extract_text_and_images(doc)
            
            # If no text was extracted, return error
            if not original_text.strip():
                return {"error": "Document is empty", "success": False}
            
            # Check total size of images
            total_image_size = sum(len(img['bytes']) for img in self._image_map.values())
            MAX_TOTAL_SIZE = 50 * 1024 * 1024  # 50MB
            
            if total_image_size > MAX_TOTAL_SIZE:
                return {
                    "error": f"Total image size ({total_image_size / 1024 / 1024:.1f}MB) exceeds maximum allowed size (50MB)",
                    "success": False
                }
            
            # Improve text using OpenAI
            messages = [
                {"role": "system", "content": self.system_prompt},
                {"role": "user", "content": original_text}
            ]
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=0.7,
                max_tokens=2048
            )
            
            improved_text = response.choices[0].message.content
            
            # Create new document with improved text
            new_doc = Document()
            
            # Process each line and apply appropriate formatting
            current_lines = []
            errors = []
            
            for line in improved_text.split('\n'):
                if not line.strip():
                    if current_lines:
                        # Process accumulated lines
                        content, style, level = self._parse_markdown_line(current_lines[0])
                        p = new_doc.add_paragraph()
                        self._apply_formatting(p, style, level)
                        
                        # Add first line with formatting
                        self._add_formatted_text(p, content)
                        
                        # Add remaining lines with the same style if any
                        for additional_line in current_lines[1:]:
                            p.add_run('\n')
                            self._add_formatted_text(p, additional_line)
                        current_lines = []
                    continue
                
                # Check if this is an image placeholder
                image_match = re.match(r'\[IMAGE:([^]]+)\]', line.strip())
                if image_match:
                    if current_lines:
                        # Process accumulated lines before adding image
                        content, style, level = self._parse_markdown_line(current_lines[0])
                        p = new_doc.add_paragraph()
                        self._apply_formatting(p, style, level)
                        
                        # Add first line with formatting
                        self._add_formatted_text(p, content)
                        
                        # Add remaining lines with the same style if any
                        for additional_line in current_lines[1:]:
                            p.add_run('\n')
                            self._add_formatted_text(p, additional_line)
                        current_lines = []
                    
                    # Add image
                    image_id = image_match.group(1).strip()
                    error = self._add_image_to_doc(new_doc, image_id)
                    if error:
                        errors.append(error)
                    
                    # Clear image from memory after adding to document
                    if image_id in self._image_map:
                        del self._image_map[image_id]
                else:
                    current_lines.append(line)
            
            # Process any remaining lines
            if current_lines:
                content, style, level = self._parse_markdown_line(current_lines[0])
                p = new_doc.add_paragraph()
                self._apply_formatting(p, style, level)
                
                # Add first line with formatting
                self._add_formatted_text(p, content)
                
                # Add remaining lines with the same style if any
                for additional_line in current_lines[1:]:
                    p.add_run('\n')
                    self._add_formatted_text(p, additional_line)
            
            return {
                "success": True,
                "original_text": original_text,
                "improvements": improved_text,
                "formatted_doc": new_doc,
                "errors": errors if errors else None
            }
            
        except Exception as e:
            logger.error(f"Error improving document: {str(e)}")
            return {"error": str(e), "success": False}

    def _apply_format_settings(self, paragraph, format_settings: Dict[str, Any]) -> None:
        """Apply format settings to a paragraph."""
        # Apply to the default run if it exists
        if paragraph.runs:
            for run in paragraph.runs:
                if 'font_size' in format_settings:
                    run.font.size = Pt(format_settings['font_size'])
                if 'font_color' in format_settings:
                    color = format_settings['font_color']
                    run.font.color.rgb = RGBColor(*color)
        
        if 'space_before' in format_settings:
            paragraph.paragraph_format.space_before = Pt(format_settings['space_before'])
        if 'space_after' in format_settings:
            paragraph.paragraph_format.space_after = Pt(format_settings['space_after'])
        if 'line_spacing' in format_settings:
            paragraph.paragraph_format.line_spacing = format_settings['line_spacing']
