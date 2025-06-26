"""Tests specifically for the image embedding functionality in DocImprover."""
import os
import pytest
import shutil
import re
import io
from pathlib import Path
from PIL import Image
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from doc_improver.document_processor import DocumentProcessor

def create_test_image(width=200, height=200, color=(255, 0, 0)):
    """Create a test image for embedding tests."""
    image = Image.new('RGB', (width, height), color=color)
    img_byte_arr = io.BytesIO()
    image.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)
    return img_byte_arr

def create_docx_with_images(output_path, num_images=2):
    """Create a test document with embedded images."""
    doc = Document()
    doc.add_heading('Test Document with Images', 0)
    doc.add_paragraph('This document contains embedded test images.')
    
    # Add multiple images
    for i in range(num_images):
        doc.add_paragraph(f'Image {i+1}:')
        img_bytes = create_test_image(color=(255, i*50, i*50))
        doc.add_picture(img_bytes)
        doc.add_paragraph(f'Text after image {i+1}')
    
    doc.save(output_path)
    return output_path

def count_images_in_docx(docx_path):
    """Count the number of embedded images in a DOCX file."""
    doc = Document(docx_path)
    image_count = 0
    
    # Check for image relationships in the main document part
    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            image_count += 1
    
    return image_count

@pytest.fixture
def processor():
    """Create a DocumentProcessor instance for testing."""
    processor = DocumentProcessor()
    yield processor
    # Clean up
    if processor._temp_dir and os.path.exists(processor._temp_dir):
        shutil.rmtree(processor._temp_dir, ignore_errors=True)

@pytest.fixture
def docx_with_images(tmp_path):
    """Create a test document with embedded images."""
    docx_path = tmp_path / "test_with_images.docx"
    create_docx_with_images(docx_path)
    return docx_path

def test_docx_to_markdown_extracts_images(processor, docx_with_images):
    """Test that _docx_to_markdown correctly extracts images from DOCX."""
    # Convert DOCX to Markdown
    markdown_str, media_path = processor._docx_to_markdown(str(docx_with_images))
    
    # Verify that media path was returned and exists
    assert media_path is not None
    assert os.path.exists(media_path)
    assert os.path.basename(media_path) == "media"
    
    # Verify that images were extracted
    assert len(os.listdir(media_path)) > 0
    
    # Verify that Markdown contains image references
    assert re.search(r'!\[.*?\]\(.*?\)', markdown_str) is not None

def test_markdown_to_docx_embeds_images(processor, docx_with_images, tmp_path):
    """Test that _markdown_to_docx correctly embeds images back into DOCX."""
    # First convert DOCX to Markdown to extract images
    markdown_str, media_path = processor._docx_to_markdown(str(docx_with_images))
    
    # Then convert Markdown back to DOCX
    output_path = tmp_path / "output.docx"
    processor._markdown_to_docx(markdown_str, str(output_path), media_path)
    
    # Verify that output DOCX exists
    assert os.path.exists(output_path)
    
    # Count images in original and output documents
    original_image_count = count_images_in_docx(docx_with_images)
    output_image_count = count_images_in_docx(output_path)
    
    # Verify that the number of images matches
    assert output_image_count == original_image_count, f"Original had {original_image_count} images, but output has {output_image_count}"

def test_end_to_end_document_improvement_with_images(processor, docx_with_images, monkeypatch):
    """Test the full document improvement process with image preservation."""
    # Mock OpenAI API client
    class MockClient:
        def __init__(self):
            self.chat = self.ChatAPI()
            
        class ChatAPI:
            def __init__(self):
                self.completions = self.CompletionsAPI()
                
            class CompletionsAPI:
                def create(self, *args, **kwargs):
                    # Extract original text from user message
                    text = next(msg['content'] for msg in kwargs['messages'] if msg['role'] == 'user')
                    # Preserve image placeholders in the response
                    placeholders = []
                    for line in text.split('\n'):
                        if '![' in line and '](' in line:
                            placeholders.append(line.strip())
                    
                    # Create improved text with preserved image placeholders
                    improved_text = ["# Improved Document", ""]
                    improved_text.extend(placeholders)
                    improved_text.append("This is the improved content.")
                    
                    # Mock response class structure
                    class Message:
                        def __init__(self, content):
                            self.content = content
                    
                    class Choice:
                        def __init__(self, message):
                            self.message = message
                    
                    class Response:
                        def __init__(self, choices):
                            self.choices = choices
                    
                    # Return properly structured response object
                    return Response([Choice(Message('\n'.join(improved_text)))])
    
    # Patch the client
    monkeypatch.setattr(processor, 'client', MockClient())
    
    # Process the document
    result = processor.improve_document(str(docx_with_images))
    
    # Verify result structure
    assert isinstance(result, dict)
    assert "markdown_path" in result
    assert "improved_docx_path" in result
    assert "media_path" in result
    assert os.path.exists(result["markdown_path"])
    assert os.path.exists(result["improved_docx_path"])
    
    # Verify that media path was captured
    assert "media_path" in result
    assert result["media_path"] is not None
    
    # Verify that improved document exists
    improved_docx_path = result["improved_docx_path"]
    assert os.path.exists(improved_docx_path)
    
    # Count images in original and improved documents
    original_image_count = count_images_in_docx(docx_with_images)
    improved_image_count = count_images_in_docx(improved_docx_path)
    
    # Verify that the number of images matches
    assert improved_image_count == original_image_count, f"Original had {original_image_count} images, but improved has {improved_image_count}"

def test_absolute_path_image_handling(processor, tmp_path):
    """Test handling of absolute paths in Markdown image references."""
    # Create test images in realistic locations
    media_dir = tmp_path / "media"
    media_dir.mkdir(exist_ok=True)
    
    # Create actual image files that can be found by pandoc
    img1 = media_dir / "image1.png"
    img2 = media_dir / "image2.jpg"
    img1.write_bytes(create_test_image().getvalue())
    img2.write_bytes(create_test_image().getvalue())
    
    # Create markdown with references to these actual paths
    markdown_text = f"""# Test Document
    
Here's an image with an file path:
![Test Image]({img1})

And another one:
![Another Image]({img2})
"""
    
    # Set processor's temp directory to our test directory to ensure images are found
    processor._temp_dir = str(tmp_path)
    
    # Convert Markdown to DOCX
    output_path = tmp_path / "absolute_path_test.docx"
    processor._markdown_to_docx(markdown_text, str(output_path), str(media_dir))
    
    # Verify output exists
    assert os.path.exists(output_path)
    
    # Check if DOCX has images
    image_count = count_images_in_docx(output_path)
    assert image_count > 0, "No images were embedded from absolute paths"
