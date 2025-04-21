"""Test configuration for DocImprover."""
import os
import pytest
import shutil
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Inches
from doc_improver.app import app as flask_app
from doc_improver.document_processor import DocumentProcessor
import io
from PIL import Image

def create_test_image(width=200, height=200):
    """Create a test image."""
    image = Image.new('RGB', (width, height), color='red')
    img_byte_arr = io.BytesIO()
    image.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)
    return img_byte_arr

@pytest.fixture
def app(tmp_path):
    """Create a test Flask application with temporary upload directory."""
    upload_dir = tmp_path / "uploads"
    upload_dir.mkdir(exist_ok=True)
    
    flask_app.config.update({
        "TESTING": True,
        "UPLOAD_FOLDER": str(upload_dir),
        "FILE_CLEANUP_AGE": timedelta(hours=1)
    })
    
    yield flask_app
    
    # Cleanup after tests
    try:
        if upload_dir.exists():
            shutil.rmtree(upload_dir, ignore_errors=True)
    except Exception:
        pass  # Ignore cleanup errors

@pytest.fixture
def client(app):
    """Create a test client."""
    return app.test_client()

@pytest.fixture
def mock_openai(monkeypatch):
    """Mock OpenAI API for testing."""
    class MockOpenAI:
        class ChatCompletion:
            class Response:
                def __init__(self, text):
                    self.choices = [type('Choice', (), {'message': type('Message', (), {'content': text})()})]
                    self.model = "gpt-4"

            @staticmethod
            def create(*args, **kwargs):
                # Extract original text from user message
                text = next(msg['content'] for msg in kwargs['messages'] if msg['role'] == 'user')
                # Preserve image placeholders in the response
                placeholders = []
                for line in text.split('\n'):
                    if '[IMAGE:' in line:
                        placeholders.append(line.strip())
                
                # Create improved text with preserved image placeholders
                improved_text = ["# Improved Document", ""]
                improved_text.extend(placeholders)
                improved_text.append("This is the improved content.")
                return MockOpenAI.ChatCompletion.Response('\n'.join(improved_text))

        def __init__(self, api_key=None, base_url=None):
            self.chat = type('Chat', (), {'completions': self.ChatCompletion})()

    monkeypatch.setattr("openai.OpenAI", MockOpenAI)
    return MockOpenAI

@pytest.fixture
def sample_doc():
    """Create a sample document for testing."""
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph.')
    doc.add_paragraph('This is another paragraph with some formatting.')
    return doc

@pytest.fixture
def sample_doc_with_images():
    """Create a sample document with embedded images for testing."""
    doc = Document()
    doc.add_heading('Test Document with Images', 0)
    doc.add_paragraph('This is a paragraph before the image.')
    
    # Add first test image
    img_bytes1 = create_test_image()
    doc.add_picture(img_bytes1, width=Inches(2.0))
    
    doc.add_paragraph('This is a paragraph between images.')
    
    # Add second test image
    img_bytes2 = create_test_image(width=300, height=300)
    doc.add_picture(img_bytes2, width=Inches(3.0))
    
    doc.add_paragraph('This is a paragraph after the image.')
    return doc

@pytest.fixture
def temp_docx(tmp_path):
    """Create a temporary Word document."""
    doc_path = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test paragraph.')
    doc.save(doc_path)
    return doc_path

@pytest.fixture
def old_session_dir(app):
    """Create an old session directory for testing cleanup."""
    session_id = "test_old_session"
    session_dir = Path(app.config['UPLOAD_FOLDER']) / session_id
    session_dir.mkdir(exist_ok=True)
    
    # Create a test file in the directory
    test_file = session_dir / "test.docx"
    test_file.write_text("test content")
    
    # Set old modification time
    old_time = datetime.now() - app.config['FILE_CLEANUP_AGE'] - timedelta(minutes=5)
    os.utime(session_dir, (old_time.timestamp(), old_time.timestamp()))
    
    return session_dir

@pytest.fixture
def recent_session_dir(app):
    """Create a recent session directory that shouldn't be cleaned up."""
    session_id = "test_recent_session"
    session_dir = Path(app.config['UPLOAD_FOLDER']) / session_id
    session_dir.mkdir(exist_ok=True)
    
    # Create a test file in the directory
    test_file = session_dir / "test.docx"
    test_file.write_text("test content")
    
    return session_dir
